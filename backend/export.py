"""Write transcripts to .txt and .docx files."""

from __future__ import annotations

from pathlib import Path

from fusion import Turn


def _format_timestamp(seconds: float) -> str:
    minutes = int(seconds // 60)
    secs = int(seconds % 60)
    return f"{minutes:02d}:{secs:02d}"


def write_txt(turns: list[Turn], path: Path) -> None:
    lines: list[str] = []
    for turn in turns:
        timestamp = _format_timestamp(turn.start)
        lines.append(f"[{timestamp}] {turn.speaker}: {turn.text}")
    path.write_text("\n\n".join(lines) + "\n", encoding="utf-8")


def write_docx(turns: list[Turn], path: Path, *, title: str | None = None) -> None:
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    if title:
        doc.add_heading(title, level=1)

    for turn in turns:
        timestamp = _format_timestamp(turn.start)
        para = doc.add_paragraph()
        speaker_run = para.add_run(f"{turn.speaker} ")
        speaker_run.bold = True
        ts_run = para.add_run(f"[{timestamp}]\n")
        ts_run.italic = True
        ts_run.font.size = Pt(9)
        para.add_run(turn.text)

    doc.save(str(path))
